import streamlit as st
import math
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import pandas as pd
from supabase import create_client, Client
from datetime import datetime, timedelta, date

# ==============================================================================
# 1. CONFIGURAÇÃO DA PÁGINA E CONEXÃO SUPABASE
# ==============================================================================
st.set_page_config(page_title="Cálculo de Energia Incidente", page_icon="⚡", layout="wide")

# 👇 CREDENCIAIS
SUPABASE_URL = "https://lfgqxphittdatzknwkqw.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImxmZ3F4cGhpdHRkYXR6a253a3F3Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzA4NzYyNzUsImV4cCI6MjA4NjQ1MjI3NX0.fZSfStTC5GdnP0Md1O0ptq8dD84zV-8cgirqIQTNO4Y"

@st.cache_resource
def init_supabase():
    try:
        return create_client(st.secrets["supabase"]["url"], st.secrets["supabase"]["key"])
    except:
        return create_client(SUPABASE_URL, SUPABASE_KEY)

try:
    supabase = init_supabase()
except Exception as e:
    st.error(f"Erro de conexão com Supabase: {e}")
    st.stop()

# ==============================================================================
# 2. FUNÇÕES AUXILIARES
# ==============================================================================
def ft(texto):
    try:
        if texto is None: return ""
        return str(texto).encode('latin-1', 'replace').decode('latin-1')
    except:
        return str(texto)

def gerar_pdf(dados):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, ft('Memorial de Cálculo - Arc Flash'), 0, 1, 'C') 
    pdf.set_font("Arial", 'I', 9)
    pdf.cell(0, 6, 'Conforme NBR 17227 / IEEE 1584', 0, 1, 'C')
    pdf.ln(4)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 6, ft(f"Local: {dados['local']}"), 0, 1, 'C')
    eq_texto = dados['eq1']
    if dados['eq2']: eq_texto += f" [{dados['eq2']}]"
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(0, 6, ft(eq_texto), 0, 1, 'C')
    pdf.ln(8)
    pdf.set_fill_color(230, 230, 230)
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(0, 7, ft("1. PARÂMETROS DE ENTRADA"), 1, 1, 'L', 1)
    pdf.set_font("Arial", size=10)
    pdf.ln(2)
    pdf.cell(95, 6, ft(f"Tensão Nominal: {dados['v']:.3f} kV"), 0, 0)
    pdf.cell(95, 6, ft(f"Corrente de Curto (Ibf): {dados['i']:.3f} kA"), 0, 1)
    pdf.cell(95, 6, ft(f"Tempo de Arco: {dados['t']:.4f} s"), 0, 0)
    pdf.cell(95, 6, ft("Configuração: VCB"), 0, 1)
    gap_txt = "(Padrao)" if dados['is_gap_std'] else "(Manual)"
    dist_txt = "(Padrao)" if dados['is_dist_std'] else "(Manual)"
    pdf.cell(95, 6, ft(f"Gap: {dados['g']:.1f} mm {gap_txt}"), 0, 0)
    pdf.cell(95, 6, ft(f"Distância: {dados['d']:.1f} mm {dist_txt}"), 0, 1)
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(0, 7, ft("2. ROTEIRO DE CÁLCULO"), 1, 1, 'L', 1)
    pdf.set_font("Courier", size=9)
    pdf.ln(2)
    pdf.cell(0, 5, f"A) Logaritmos:", 0, 1)
    pdf.cell(0, 5, f"   Log(Ibf)={math.log10(dados['i']):.4f} | Log(Gap)={math.log10(dados['g']):.4f}", 0, 1)
    pdf.ln(2)
    pdf.cell(0, 5, ft(f"B) Energia Base (En):"), 0, 1)
    pdf.cell(0, 5, f"   Log(En) = {dados['lg_en']:.4f} -> En = {dados['en_base']:.4f} cal/cm2", 0, 1)
    pdf.ln(2)
    pdf.cell(0, 5, ft(f"C) Fatores:"), 0, 1)
    pdf.cell(0, 5, f"   Tempo ({dados['t']}s/0.2s): {dados['fator_t']:.2f}", 0, 1)
    pdf.cell(0, 5, f"   Distancia (610/{dados['d']})^2: {dados['fator_d']:.3f}", 0, 1)
    pdf.cell(0, 5, f"   Fator Tensao: {dados['fator_v']}", 0, 1)
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(0, 7, ft("3. RESULTADO E CLASSIFICAÇÃO"), 1, 1, 'L', 1)
    pdf.ln(3)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, ft(f"Energia Incidente: {dados['e']:.2f} cal/cm²"), 0, 1)
    pdf.set_font("Arial", size=11)
    pdf.set_text_color(0, 0, 0)
    if dados['e'] > 40: pdf.set_text_color(200, 0, 0)
    elif dados['e'] >= 8: pdf.set_text_color(200, 100, 0)
    pdf.cell(0, 8, ft(f"Classificação: {dados['cat']}"), 0, 1)
    pdf.set_text_color(0, 0, 0) 
    pdf.ln(5)
    pdf.set_font("Arial", 'I', 8)
    pdf.cell(0, 5, ft("Nota: A vestimenta deve possuir ATPV superior à energia calculada."), 0, 1)
    return pdf.output(dest='S').encode('latin-1')

def gerar_word(dados):
    doc = Document()
    head = doc.add_heading('Memorial - Arc Flash', 0) 
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_local = doc.add_paragraph()
    p_local.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p_local.add_run(f"Local: {dados['local']}")
    run_l.bold = True
    run_l.font.size = Pt(12)
    eq_texto = dados['eq1']
    if dados['eq2']: eq_texto += f" [{dados['eq2']}]"
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_eq = p_eq.add_run(eq_texto)
    run_eq.bold = True
    run_eq.font.size = Pt(11)
    doc.add_paragraph("-" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('1. Parâmetros', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Tensão: {dados['v']:.3f} kV | Corrente: {dados['i']:.3f} kA | Tempo: {dados['t']:.4f} s\n")
    p.add_run(f"Gap: {dados['g']:.1f} mm | Distância: {dados['d']:.1f} mm\n")
    p.add_run("Configuração: VCB")
    doc.add_heading('2. Resultado', level=1)
    p_res = doc.add_paragraph()
    run_res = p_res.add_run(f"{dados['e']:.2f} cal/cm²")
    run_res.bold = True
    run_res.font.size = Pt(16)
    doc.add_paragraph(f"Classificação: {dados['cat']}")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==============================================================================
# 3. LÓGICA DE LOGIN / CADASTRO
# ==============================================================================
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
    st.session_state['user_role'] = None
    st.session_state['user_login'] = None
    st.session_state['user_name'] = None

if not st.session_state['logged_in']:
    col1, col2, col3 = st.columns([1,1,1])
    with col2:
        st.title("🔒 Sistema Arc Flash")
        modo = st.radio("Acesso:", ["Entrar", "Criar Conta"], horizontal=True)
        
        if modo == "Entrar":
            with st.form("login_form"):
                nome_login = st.text_input("Nome", key="input_nome_login_v2") 
                pwd = st.text_input("Senha", type="password", key="input_pwd_login_v2")
                submitted = st.form_submit_button("Entrar", type="primary")
                
                if submitted:
                    try:
                        res = supabase.table('users').select("*").or_(f"name.eq.{nome_login},username.eq.{nome_login}").eq('password', pwd).execute()
                        
                        if res.data:
                            data = res.data[0]
                            # 1. Verifica se está aprovado (Ativo/Bloqueado)
                            if data.get('approved'):
                                
                                # 2. VERIFICAÇÃO DE CONTRATO (VALIDADE)
                                exp_str = data.get('expiration_date')
                                contrato_valido = True
                                
                                if exp_str:
                                    try:
                                        # Converte string do banco para data
                                        validade = datetime.strptime(exp_str, '%Y-%m-%d').date()
                                        if date.today() > validade:
                                            contrato_valido = False
                                    except: pass # Se data inválida, ignora e libera (ou bloqueia, depende da regra)

                                if contrato_valido:
                                    st.session_state['logged_in'] = True
                                    
                                    email_db = data.get('username')
                                    if email_db == 'admin' or nome_login.lower() == 'admin':
                                        st.session_state['user_role'] = 'admin'
                                    else:
                                        st.session_state['user_role'] = 'user'
                                    
                                    st.session_state['user_name'] = data.get('name')
                                    st.session_state['user_login'] = email_db
                                    
                                    # LOG AUTOMÁTICO
                                    try:
                                        supabase.table("arc_flash_history").insert({
                                            "username": email_db,
                                            "tag_equipamento": "🟢 LOGIN DO SISTEMA",
                                            "tensao_kv": 0.0,
                                            "corrente_ka": 0.0,
                                            "tempo_s": 0.0,
                                            "distancia_mm": 0.0,
                                            "energia_cal": 0.0
                                        }).execute()
                                    except: pass
                                    
                                    st.rerun()
                                else:
                                    st.error(f"⛔ Seu contrato de licença expirou em {validade.strftime('%d/%m/%Y')}. Contate o administrador.")
                            else:
                                st.warning("🚫 Usuário pendente de aprovação ou bloqueado pelo Administrador.")
                        else:
                            st.error("Nome ou senha incorretos.")
                    except Exception as e:
                        st.error(f"Erro de conexão: {e}")
        
        else: # Criar Conta
            with st.form("cadastro_form"):
                st.markdown("### Novo Cadastro")
                new_name = st.text_input("Nome")
                new_email = st.text_input("E-mail")
                new_pass = st.text_input("Defina sua Senha", type="password")
                
                reg_btn = st.form_submit_button("Solicitar Acesso")
                
                if reg_btn:
                    if new_email and new_name and new_pass:
                        try:
                            check = supabase.table('users').select("*").eq('username', new_email).execute()
                            if check.data:
                                st.error("Este e-mail já está cadastrado.")
                            else:
                                payload = {
                                    "username": new_email, 
                                    "name": new_name,
                                    "password": new_pass,
                                    "approved": False
                                    # expiration_date será setado quando o Admin aprovar
                                }
                                supabase.table('users').insert(payload).execute()
                                st.success("✅ Cadastro realizado! Aguarde aprovação.")
                        except Exception as e:
                            st.error(f"Erro ao cadastrar: {e}")
                    else:
                        st.warning("Preencha todos os campos.")
    st.stop()

# ==============================================================================
# 4. APP PRINCIPAL
# ==============================================================================

st.sidebar.success(f"Olá, {st.session_state['user_name']}")

# 4.1 Troca de Senha
with st.sidebar.expander("🔑 Alterar Minha Senha"):
    with st.form("mudar_senha"):
        senha_atual = st.text_input("Senha Atual", type="password")
        nova_senha = st.text_input("Nova Senha", type="password")
        btn_mudar = st.form_submit_button("Atualizar")
        if btn_mudar:
            try:
                chk = supabase.table('users').select("*").eq('username', st.session_state['user_login']).eq('password', senha_atual).execute()
                if chk.data:
                    supabase.table('users').update({'password': nova_senha}).eq('username', st.session_state['user_login']).execute()
                    st.success("Senha alterada!")
                else:
                    st.error("Senha atual incorreta.")
            except: pass

# 4.2 Painel Admin
isAdmin = (st.session_state['user_role'] == 'admin')
if isAdmin:
    st.sidebar.markdown("---")
    st.sidebar.subheader("🛡️ Admin Panel")
    
    # -----------------------------------------------
    # APROVAÇÃO E RENOVAÇÃO DE LICENÇA (1 ANO)
    # -----------------------------------------------
    try:
        # Busca aprovados=False (pendentes ou bloqueados)
        pendentes = supabase.table('users').select("*").eq('approved', False).execute()
        if pendentes.data:
            st.sidebar.warning(f"Pendentes/Bloqueados: {len(pendentes.data)}")
            
            lista_pend = {f"{u['name']} ({u['username']})": u['username'] for u in pendentes.data}
            sel_display = st.sidebar.selectbox("Liberar/Renovar:", list(lista_pend.keys()))
            sel_email = lista_pend[sel_display]
            
            if st.sidebar.button(f"✅ Liberar + 1 Ano"):
                # Calcula expiração: Hoje + 365 dias
                validade_nova = (datetime.now() + timedelta(days=365)).strftime('%Y-%m-%d')
                
                supabase.table('users').update({
                    'approved': True,
                    'expiration_date': validade_nova
                }).eq('username', sel_email).execute()
                
                st.sidebar.success(f"Acesso liberado até {datetime.strptime(validade_nova, '%Y-%m-%d').strftime('%d/%m/%Y')}!")
                st.rerun()
        else:
            st.sidebar.info("Sem aprovações pendentes.")
    except Exception as e:
        st.sidebar.error(f"Erro ao carregar pendentes. Verifique se criou a coluna expiration_date.")
    
    st.sidebar.markdown("---")
    
    # -----------------------------------------------
    # BLOQUEIO DE USUÁRIO (SEM EXCLUIR)
    # -----------------------------------------------
    try:
        # Busca usuários ativos (approved=True) exceto admin e o próprio
        active_users = supabase.table('users').select("*").eq('approved', True).neq('username', 'admin').neq('username', st.session_state['user_login']).execute()
        
        if active_users.data:
            lista_ativos = {f"{u['name']} ({u['username']})": u['username'] for u in active_users.data}
            
            user_block_display = st.sidebar.selectbox("Bloquear Acesso:", ["..."] + list(lista_ativos.keys()))
            
            if user_block_display != "...":
                email_block = lista_ativos[user_block_display]
                if st.sidebar.button(f"🚫 Bloquear {email_block}"):
                    supabase.table('users').update({'approved': False}).eq('username', email_block).execute()
                    st.sidebar.warning(f"Usuário {email_block} foi bloqueado.")
                    st.rerun()
    except: pass

    st.sidebar.markdown("---")
    
    # -----------------------------------------------
    # EXCLUSÃO DEFINITIVA
    # -----------------------------------------------
    try:
        all_users = supabase.table('users').select("*").neq('username', 'admin').neq('username', st.session_state['user_login']).execute()
        if all_users.data:
            users_map = {f"{u['name']} ({u['username']})": u['username'] for u in all_users.data}
            user_del_display = st.sidebar.selectbox("Excluir Definitivamente:", ["..."] + list(users_map.keys()))
            
            if user_del_display != "...":
                email_del = users_map[user_del_display]
                if st.sidebar.button(f"🗑️ Confirmar Exclusão"):
                    try:
                        supabase.table('arc_flash_history').delete().eq('username', email_del).execute()
                        supabase.table('users').delete().eq('username', email_del).execute()
                        st.sidebar.success("Usuário e histórico excluídos.")
                        st.rerun()
                    except:
                        st.sidebar.error("Erro na exclusão.")
    except: pass

st.sidebar.markdown("---")
if st.sidebar.button("Sair do Sistema"):
    st.session_state['logged_in'] = False
    st.rerun()

# --- ÁREA CENTRAL ---
st.markdown(f"### ⚡ Cálculo de Arc Flash")

if 'corrente_stored' not in st.session_state: st.session_state['corrente_stored'] = 17.0
if 'resultado_icc_detalhe' not in st.session_state: st.session_state['resultado_icc_detalhe'] = None
if 'ultimo_calculo' not in st.session_state: st.session_state['ultimo_calculo'] = None

# ABAS
if isAdmin:
    tab1, tab2, tab3 = st.tabs(["🔥 Energia Incidente", "🧮 Icc (Curto)", "📂 Histórico (Admin)"])
else:
    tab1, tab2 = st.tabs(["🔥 Energia Incidente", "🧮 Icc (Curto)"])
    tab3 = None

# === ABA 1: CÁLCULO ===
with tab1:
    st.subheader("Análise de Energia")
    with st.container(border=True):
        st.caption("Identificação")
        local_input = st.text_input("Local", placeholder="Ex: Sala Elétrica 01")
        c_eq1, c_eq2 = st.columns(2)
        with c_eq1: eq1_input = st.text_input("Equipamento", placeholder="Ex: QGBT Geral")
        with c_eq2: eq2_input = st.text_input("Detalhe", placeholder="Ex: Disjuntor Entrada")

    st.write("")
    st.info("Parâmetros do Arco:")
    c1, c2, c3 = st.columns(3)
    with c1: tensao = st.number_input("1. Tensão (kV)", value=13.80, format="%.3f")
    with c2: corrente = st.number_input("2. Corrente (kA)", key="corrente_stored", format="%.3f")
    with c3: tempo = st.number_input("3. Tempo (s)", value=0.500, format="%.4f")

    c4, c5 = st.columns(2)
    with c4: gap = st.number_input("Gap (mm)", value=0.0, step=1.0)
    with c5: distancia = st.number_input("Distância (mm)", value=0.0, step=10.0)

    def calcular_completo():
        g_c = gap if gap > 0 else (152.0 if tensao >= 1.0 else 25.0)
        d_c = distancia if distancia > 0 else (914.0 if tensao >= 1.0 else 457.2)
        lg_i = math.log10(corrente) if corrente > 0 else 0
        if tensao < 1.0:
            k_base, k_i, k_g = -0.555, 1.081, 0.0011
            x_dist = 2.0
            fator_v = 0.85 if tensao < 0.6 else 1.0
        else:
            k_base, k_i, k_g = -0.555, 1.081, 0.0011
            x_dist = 2.0
            fator_v = 1.15
        lg_en = k_base + (k_i * lg_i) + (k_g * g_c)
        en_base = 10 ** lg_en
        fator_t = tempo / 0.2
        fator_d = (610 / d_c) ** x_dist
        e_final = 1.0 * en_base * fator_t * fator_d * fator_v
        
        if e_final < 1.2: cat, cor = "Risco Mínimo", "green"
        elif e_final < 4.0: cat, cor = "Cat 1 / 2", "orange"
        elif e_final < 8.0: cat, cor = "Cat 2", "darkorange"
        elif e_final < 40.0: cat, cor = "Cat 3 / 4", "red"
        else: cat, cor = "PERIGO", "black"

        return {
            'local': local_input, 'eq1': eq1_input, 'eq2': eq2_input,
            'v': tensao, 'i': corrente, 't': tempo, 'g': g_c, 'd': d_c,
            'lg_en': lg_en, 'en_base': en_base,
            'fator_t': fator_t, 'fator_d': fator_d, 'fator_v': fator_v, 'is_gap_std': gap<=0, 'is_dist_std': distancia<=0,
            'e': e_final, 'cat': cat, 'cor': cor
        }

    if st.button("CALCULAR", type="primary", use_container_width=True):
        if tensao > 0 and corrente > 0:
            st.session_state['ultimo_calculo'] = calcular_completo()
        else:
            st.warning("Preencha os dados.")

    if st.session_state['ultimo_calculo']:
        res = st.session_state['ultimo_calculo']
        st.divider()
        st.markdown(f"**Resultado:** {res['local']} - {res['eq1']}")
        cr1, cr2 = st.columns([1, 2])
        cr1.metric("Energia Incidente", f"{res['e']:.2f} cal/cm²")
        cr2.markdown(f"<div style='background-color:{res['cor']};color:white;padding:15px;text-align:center;border-radius:10px;'><h3>{res['cat']}</h3></div>", unsafe_allow_html=True)
        
        st.divider()
        st.caption("Ações:")
        
        cb1, cb2, cb3 = st.columns(3)
        with cb1:
            if st.button("💾 Gravar no Banco de Dados", use_container_width=True):
                try:
                    payload = {
                        "username": st.session_state['user_login'],
                        "tag_equipamento": res['eq1'] if res['eq1'] else "Sem Tag",
                        "tensao_kv": res['v'], "corrente_ka": res['i'], "tempo_s": res['t'], "distancia_mm": res['d'],
                        "energia_cal": float(f"{res['e']:.2f}")
                    }
                    supabase.table("arc_flash_history").insert(payload).execute()
                    st.toast("✅ Salvo com sucesso!", icon="💾")
                except Exception as e:
                    st.error(f"Erro: {e}")
        with cb2:
            st.download_button("📥 Baixar PDF", data=gerar_pdf(res), file_name="memorial.pdf", mime="application/pdf", use_container_width=True)
        with cb3:
            st.download_button("📝 Baixar Word", data=gerar_word(res), file_name="memorial.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

# === ABA 2: ICC ===
with tab2:
    st.subheader("Estimativa Curto-Circuito")
    def atualizar_icc():
        try:
            t_kva = st.session_state['k_kva']
            t_v = st.session_state['k_v']
            t_z = st.session_state['k_z']
            if t_v > 0 and t_z > 0:
                i_nom = (t_kva * 1000) / (math.sqrt(3) * t_v)
                i_cc = i_nom / (t_z / 100)
                i_mot = 4 * i_nom if st.session_state['k_motor'] else 0
                total = (i_cc + i_mot) / 1000
                st.session_state['corrente_stored'] = total
                st.session_state['resultado_icc_detalhe'] = {'total': total}
                st.toast(f"Calculado: {total:.3f} kA", icon="✅")
        except: pass

    c1, c2 = st.columns(2)
    with c1:
        st.number_input("Potência Trafo (kVA)", value=1000.0, step=100.0, key="k_kva")
        st.number_input("Tensão Sec. (V)", value=380.0, step=10.0, key="k_v")
    with c2:
        st.number_input("Impedância Z (%)", value=5.0, step=0.1, key="k_z")
        st.checkbox("Considerar Motores?", value=True, key="k_motor")
    st.write("")
    st.button("Calcular Icc", on_click=atualizar_icc, type="primary", use_container_width=True)
    if st.session_state['resultado_icc_detalhe']:
        st.divider()
        st.metric("Icc Estimada", f"{st.session_state['resultado_icc_detalhe']['total']:.3f} kA")
        st.success("Valor copiado automaticamente para a Aba 1.")

# === ABA 3: HISTÓRICO (ADMIN) ===
if isAdmin and tab3:
    with tab3:
        st.header("📂 Banco de Dados de Simulações")
        if st.button("🔄 Atualizar Dados"):
            st.rerun()
        try:
            res_hist = supabase.table("arc_flash_history").select("*").order("created_at", desc=True).execute()
            if res_hist.data:
                df = pd.DataFrame(res_hist.data)
                cols_map = {
                    'created_at': 'Data/Hora', 'username': 'Usuário', 'tag_equipamento': 'Tag',
                    'energia_cal': 'Energia', 'tensao_kv': 'kV', 'corrente_ka': 'kA'
                }
                df.rename(columns=cols_map, inplace=True)
                # UTC -> SP
                if 'Data/Hora' in df.columns:
                    df['Data/Hora'] = pd.to_datetime(df['Data/Hora'], utc=True)
                    df['Data/Hora'] = df['Data/Hora'].dt.tz_convert('America/Sao_Paulo').dt.strftime('%d/%m/%Y %H:%M')
                final_cols = [c for c in cols_map.values() if c in df.columns]
                st.dataframe(df[final_cols], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhum registro encontrado.")
        except Exception as e:
            st.error(f"Erro ao carregar banco: {e}")
